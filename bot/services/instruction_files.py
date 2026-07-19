from __future__ import annotations

import logging
from pathlib import Path

logger = logging.getLogger(__name__)

# Telegram extensions we can read as instruction text.
SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt", ".md"}

# Cap on accumulated material text stored on a catalog exercise. Raw materials are now
# read once (during the AI understanding check) rather than resent in every analysis
# batch — the batches get the compact understanding brief instead — so this can be
# generous. Measured: ~0.25 ₽ per 1k input tokens, and 60k chars ≈ 18k tokens read once.
MAX_INSTRUCTIONS_CHARS = 60000


class InstructionExtractionError(RuntimeError):
    pass


def is_supported_instruction(file_name: str | None) -> bool:
    return Path(file_name or "").suffix.lower() in SUPPORTED_EXTENSIONS


def extract_instruction_text(path: Path) -> str:
    """Extract plain text from an instruction file (.pdf/.docx/.txt/.md).

    Raises InstructionExtractionError for unsupported types or unreadable files.
    """
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        text = _extract_pdf(path)
    elif suffix == ".docx":
        text = _extract_docx(path)
    elif suffix in {".txt", ".md"}:
        text = path.read_text(encoding="utf-8", errors="ignore")
    else:
        raise InstructionExtractionError(f"Неподдерживаемый формат инструкции: {suffix or '—'}")

    cleaned = _normalize(text)
    if not cleaned:
        raise InstructionExtractionError("Из файла не удалось извлечь текст (возможно, это скан без текста).")
    return cleaned


def combine_instructions(parts: list[tuple[str, str]]) -> tuple[str, bool]:
    """Build the whole material text from scratch, tagging each chunk with its filename.

    Returns (text, truncated). Used when the set of materials changes (a file is
    removed), where appending is not enough — the text must be rebuilt from what is
    left, otherwise a deleted file would keep influencing the analysis.
    """
    blocks = [f"=== {source} ===\n{text}".strip() for source, text in parts if (text or "").strip()]
    combined = "\n\n".join(blocks)
    if len(combined) > MAX_INSTRUCTIONS_CHARS:
        return combined[:MAX_INSTRUCTIONS_CHARS].rstrip(), True
    return combined, False


def append_instructions(existing: str | None, addition: str, *, source: str | None = None) -> str:
    """Concatenate a new instruction chunk onto the exercise's accumulated text.

    Keeps the total under MAX_INSTRUCTIONS_CHARS and tags each chunk with its filename.
    """
    header = f"=== {source} ===\n" if source else ""
    block = f"{header}{addition}".strip()
    combined = f"{existing.strip()}\n\n{block}" if (existing or "").strip() else block
    if len(combined) > MAX_INSTRUCTIONS_CHARS:
        combined = combined[:MAX_INSTRUCTIONS_CHARS].rstrip()
    return combined


def _extract_pdf(path: Path) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as exc:  # pragma: no cover - dependency guard
        raise InstructionExtractionError("Чтение PDF недоступно: не установлен pypdf.") from exc

    try:
        reader = PdfReader(str(path))
        pages = [page.extract_text() or "" for page in reader.pages]
    except Exception as exc:  # pypdf raises a variety of errors on malformed files
        logger.exception("Failed to read instruction PDF %s", path)
        raise InstructionExtractionError("Не удалось прочитать PDF-файл инструкции.") from exc
    return "\n".join(pages)


def _extract_docx(path: Path) -> str:
    try:
        from docx import Document
    except ImportError as exc:  # pragma: no cover - dependency guard
        raise InstructionExtractionError("Чтение DOCX недоступно: не установлен python-docx.") from exc

    try:
        document = Document(str(path))
    except Exception as exc:
        logger.exception("Failed to read instruction DOCX %s", path)
        raise InstructionExtractionError("Не удалось прочитать DOCX-файл инструкции.") from exc

    lines = [para.text for para in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                lines.append(" | ".join(cells))
    return "\n".join(lines)


def _normalize(text: str) -> str:
    lines = [line.strip() for line in (text or "").splitlines()]
    return "\n".join(line for line in lines if line).strip()
