from __future__ import annotations

from collections import defaultdict
from pathlib import Path
import re
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor


LEVEL_DESCRIPTIONS = {
    3: "Демонстрирует всегда. Компетенция проявляется на превосходном уровне и является очевидной сильной стороной.",
    2.5: "Демонстрирует почти всегда. Компетенция выражена выше достаточного уровня, отдельные индикаторы могут проявляться неустойчиво.",
    2: "Демонстрирует в большинстве случаев. Компетенция выражена на достаточном уровне, отдельные элементы требуют развития.",
    1.5: "Демонстрирует ситуативно. Поведение неустойчиво, рекомендуется планомерное развитие.",
    1: "Частично демонстрирует. Компетенция является зоной роста, проявления нестабильны.",
    0.5: "Низкий уровень. Есть единичные позитивные проявления.",
    0: "Не демонстрирует. Отсутствует проявление компетенции или преобладает противоположное поведение.",
}


def build_participant_report_text(
    *,
    participant_name: str,
    exercise_results: list[dict[str, Any]],
) -> tuple[str, dict[str, Any]]:
    summary = _aggregate(exercise_results)
    lines = [
        f"Индивидуальный отчет участника ассессмент-центра",
        f"Участник: {participant_name}",
        "",
        "Краткий вывод",
        "Отчет сформирован автоматически на основе заполненных блокнотов наблюдателя.",
        "",
    ]

    for competence, data in summary.items():
        lines.append(f"Компетенция: {competence}")
        lines.append(f"Средний уровень: {data['avg_level']}")
        lines.append("Сильные стороны:")
        lines.extend(_bullet_lines(data["strengths"]))
        lines.append("Зоны роста:")
        lines.extend(_bullet_lines(data["growth_zones"]))
        lines.append("Рекомендации:")
        lines.extend(_bullet_lines(_recommendations(data["growth_zones"])))
        lines.append("")

    return "\n".join(lines).strip(), {"competencies": summary}


def build_development_plan_text(
    *,
    participant_name: str,
    report_json: dict[str, Any],
) -> tuple[str, dict[str, Any]]:
    competencies = report_json.get("competencies", {})
    lines = [
        "Индивидуальный план развития",
        f"Участник: {participant_name}",
        "",
        "Цель ИПР",
        "Сфокусироваться на индикаторах, которые не проявились в упражнениях ассессмент-центра.",
        "",
    ]

    plan: dict[str, Any] = {}
    for competence, data in competencies.items():
        growth_zones = data.get("growth_zones", [])
        if not growth_zones:
            continue
        ipr = data.get("ipr") or {}
        resources = [_format_course(c) for c in (data.get("courses") or [])]
        resources += list(data.get("literature") or [])
        plan[competence] = {"ipr": ipr, "resources": resources}
        lines.append(f"Компетенция: {competence}")
        lines.append(f"Цель развития: {ipr.get('goal', '')}")
        lines.append("")

    if not plan:
        lines.append("Явных зон роста по обработанным упражнениям не найдено.")

    return "\n".join(lines).strip(), {"plan": plan}


def save_text_report(path: Path, text: str) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(text, encoding="utf-8")


def save_participant_report_docx(
    *,
    path: Path,
    participant_name: str,
    center_name: str | None,
    exercise_names: list[str],
    report_json: dict[str, Any],
) -> None:
    doc = Document()
    _setup_document(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("ОТЧЕТ ПО РЕЗУЛЬТАТАМ\nОЦЕНКИ КОМПЕТЕНЦИЙ")
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(31, 78, 121)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run(f"Участник: {participant_name}").bold = True
    if center_name:
        center_line = doc.add_paragraph()
        center_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
        center_line.add_run(f"Ассессмент-центр: {center_name}")

    _add_heading(doc, "Оценочные процедуры")
    if exercise_names:
        for name in exercise_names:
            _add_bullet(doc, name)
    else:
        doc.add_paragraph("Данные по упражнениям не указаны.")

    _add_heading(doc, "Шкала оценки")
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "Уровень"
    table.rows[0].cells[1].text = "Описание"
    for level in [3, 2.5, 2, 1.5, 1, 0.5, 0]:
        row = table.add_row().cells
        row[0].text = str(level).replace(".", ",")
        row[1].text = LEVEL_DESCRIPTIONS[level]

    _add_heading(doc, "Результаты центра оценки и развития")
    competencies = report_json.get("competencies", {})
    result_table = doc.add_table(rows=1, cols=2)
    result_table.style = "Table Grid"
    result_table.rows[0].cells[0].text = "Компетенция"
    result_table.rows[0].cells[1].text = "Средний уровень"
    for competence, data in competencies.items():
        row = result_table.add_row().cells
        row[0].text = str(competence)
        row[1].text = str(data.get("avg_level", 0)).replace(".", ",")

    for competence, data in competencies.items():
        doc.add_page_break()
        _add_heading(doc, str(competence))
        level = float(data.get("avg_level", 0) or 0)
        doc.add_paragraph(f"Ваш результат: {str(level).replace('.', ',')}")
        doc.add_paragraph(_nearest_level_description(level))

        _add_heading(doc, "Сильные стороны", level=2)
        strengths = data.get("report_strengths") or data.get("strengths", [])
        _add_list_or_empty(doc, strengths, "Не выявлено по обработанным упражнениям.")

        _add_heading(doc, "Зоны роста", level=2)
        growth_zones = data.get("report_growth_zones") or data.get("growth_zones", [])
        _add_list_or_empty(doc, growth_zones, "Не выявлено по обработанным упражнениям.")

        _add_heading(doc, "Рекомендации по развитию", level=2)
        recommendations = data.get("recommendations") or _recommendations(data.get("growth_zones", []))
        _add_list_or_empty(doc, recommendations, "Продолжать закреплять сильные стороны.")

        literature = data.get("literature") or []
        if literature:
            _add_heading(doc, "Литература", level=2)
            for item in literature:
                _add_bullet(doc, item)

        courses = data.get("courses") or []
        if courses:
            _add_heading(doc, "Тренинговые блоки", level=2)
            for course in courses:
                _add_bullet(doc, _format_course(course))

        practice_tasks = data.get("practice_tasks") or []
        if practice_tasks:
            _add_heading(doc, "Задания для закрепления", level=2)
            for task in practice_tasks:
                _add_bullet(doc, task)

    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)


# Fixed development-action categories of the customer's IPR template (70/20/10).
_IPR_CATEGORIES = [
    ("workplace", "Развитие на рабочем месте",
     "конкретные поручения/задачи из ежедневной работы, которые способствуют развитию компетенции"),
    ("projects", "Специальные задачи и проекты",
     "участие в проекте или временное назначение, требующие более высокого уровня компетенции"),
    ("feedback", "Поиск обратной связи",
     "обсуждение с коллегами, подчинёнными, руководителем своей работы с точки зрения данной компетенции"),
    ("mentoring", "Наставничество",
     "наблюдение за человеком, у которого данная компетенция развита высоко, обсуждение его опыта"),
    ("self_study", "Самообучение",
     "анализ своей работы, чтение литературы, прохождение курсов, вебинаров, семинаров"),
]


def save_development_plan_docx(
    *,
    path: Path,
    participant_name: str,
    center_name: str | None,
    plan_json: dict[str, Any],
) -> None:
    doc = Document()
    _setup_document(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("ИНДИВИДУАЛЬНЫЙ ПЛАН РАЗВИТИЯ")
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(31, 78, 121)

    _add_heading(doc, "Что такое ИПР")
    doc.add_paragraph(
        "Индивидуальный план развития (ИПР) — документ, в котором прописаны долгосрочные и "
        "краткосрочные карьерные цели сотрудника, его области для развития, а также подробный план "
        "действий и ресурсы, необходимые для достижения поставленных целей."
    )
    _add_heading(doc, "Почему важно составлять ИПР", level=2)
    for item in (
        "Достижение целей организации: ИПР согласует цели сотрудника и организации.",
        "Повышение производительности: новые знания и навыки повышают эффективность.",
        "Повышение мотивации: постановка и достижение целей повышает вовлечённость.",
        "Регулярная обратная связь: ИПР способствует взаимодействию руководителя и сотрудника.",
        "Эффективное распределение ресурсов: развитие направляется в зоны наибольшего эффекта.",
        "Организационная гибкость: новые навыки расширяют возможности сотрудника.",
    ):
        _add_bullet(doc, item)

    _add_heading(doc, "Модель 70/20/10", level=2)
    _add_bullet(doc, "70% — обучение на рабочем месте: применение знаний в реальных задачах.")
    _add_bullet(doc, "20% — опыт других людей: обратная связь, наставничество, наблюдение.")
    _add_bullet(doc, "10% — формальное обучение: курсы, книги, вебинары, тренинги.")

    _add_heading(doc, "Шаблон индивидуального плана развития")
    _ipr_header_table(doc, participant_name)

    plan = plan_json.get("plan", {})
    if not plan:
        doc.add_paragraph("Явных зон роста по обработанным упражнениям не найдено.")
    for index, (competence, data) in enumerate(plan.items(), start=1):
        ipr = data.get("ipr") or {}
        resources = data.get("resources") or []
        doc.add_paragraph()
        _add_heading(doc, f"Область развития №{index}", level=2)
        doc.add_paragraph(f"Компетенция: {competence}")
        doc.add_paragraph(f"Цель развития: {ipr.get('goal', '')}")
        _ipr_actions_table(doc, ipr=ipr, resources=resources)

    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)


def _ipr_header_table(doc: Document, participant_name: str) -> None:
    rows = [
        ("ФИО сотрудника", participant_name),
        ("Должность сотрудника", ""),
        ("ФИО руководителя", ""),
        ("Должность руководителя", ""),
        ("Период выполнения ИПР", ""),
        ("Дата заполнения", ""),
    ]
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    for ri, (label, value) in enumerate(rows):
        table.rows[ri].cells[0].text = label
        table.rows[ri].cells[1].text = value


def _ipr_actions_table(doc: Document, *, ipr: dict[str, Any], resources: list[str]) -> None:
    headers = ["Действия", "Ресурсы", "Сроки", "Ожидаемые результаты", "Оценка о выполнении"]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for ci, head in enumerate(headers):
        cell = table.rows[0].cells[ci]
        cell.text = head
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True

    expected = str(ipr.get("expected_results") or "")
    for ri, (key, title, description) in enumerate(_IPR_CATEGORIES):
        cells = table.add_row().cells
        action = str(ipr.get(key) or "")
        action_text = f"{title}\n({description})"
        if action:
            action_text += f"\n• {action}"
        cells[0].text = action_text
        if key == "self_study" and resources:
            cells[1].text = "\n".join(f"• {r}" for r in resources)
        cells[3].text = expected if ri == 0 else ""


def _aggregate(exercise_results: list[dict[str, Any]]) -> dict[str, dict[str, Any]]:
    by_competence: dict[str, dict[str, Any]] = defaultdict(
        lambda: {"levels": [], "indicator_statuses": defaultdict(list), "indicator_quotes": defaultdict(list)}
    )

    for result in exercise_results:
        levels = result.get("levels", {})
        for competence, level_data in levels.items():
            if isinstance(level_data, dict):
                by_competence[competence]["levels"].append(float(level_data.get("level") or 0))

        indicators = {
            item.get("indicator_id"): item
            for item in result.get("indicators", [])
            if isinstance(item, dict)
        }
        for item in result.get("results", []):
            if not isinstance(item, dict):
                continue
            indicator = indicators.get(item.get("indicator_id"), {})
            competence = str(indicator.get("competence") or "Компетенция не указана")
            indicator_text = str(indicator.get("indicator") or item.get("indicator_id"))
            status = item.get("status")
            if status == "НЗ":
                continue
            by_competence[competence]["indicator_statuses"][indicator_text].append(status)
            for evidence in item.get("evidence", []) or []:
                if isinstance(evidence, dict) and evidence.get("quote"):
                    by_competence[competence]["indicator_quotes"][indicator_text].append(str(evidence["quote"]))

    output: dict[str, dict[str, Any]] = {}
    for competence, data in by_competence.items():
        strengths: list[str] = []
        growth_zones: list[str] = []
        for indicator_text, statuses in data["indicator_statuses"].items():
            clean_text = clean_indicator_text(indicator_text)
            if statuses and all(status == "+" for status in statuses):
                strengths.append(clean_text)
            elif any(status == "-" for status in statuses):
                growth_zones.append(clean_text)

        levels = data["levels"]
        avg_level = round(sum(levels) / len(levels), 2) if levels else 0
        output[competence] = {
            "avg_level": avg_level,
            "strengths": strengths,
            "growth_zones": growth_zones,
        }
    return output


# Assessor-facing clutter baked into notebook indicator texts that should not appear
# in a participant-facing report.
# Removed as a leading prefix only:
_PREFIX_PATTERNS = [
    re.compile(r"^\s*доп(?:олнительный)?\.?\s*замер\s*:?", re.IGNORECASE),
]
# Markers that introduce assessor examples/instructions: cut from the marker to the end.
_CUT_TO_END_PATTERNS = [
    re.compile(r"\bставим\s+замер.*", re.IGNORECASE),
    re.compile(r"\bсмотр(?:им|ите)\b.*", re.IGNORECASE),
    re.compile(r"\bуточн(?:ить|им|ите)\b[^.]*интервью.*", re.IGNORECASE),
    re.compile(r"\bнапример\b.*", re.IGNORECASE),
    re.compile(r"(?:^|\s)ИЛИ\s.*"),
]
# Removed wherever they occur:
_TOKEN_PATTERNS = [
    re.compile(r"\bдоп(?:олнительный)?\.?\s*замер\s*:?", re.IGNORECASE),
    re.compile(r"(?<!\w)не\s+закрываем", re.IGNORECASE),
    re.compile(r"(?<!\w)закрываем", re.IGNORECASE),
    re.compile(r"\b(?:т\.е\.|и\s*т\.\s*д\.?|и\s*т\.\s*п\.?)\b[,.]?", re.IGNORECASE),
]


def clean_indicator_text(text: str) -> str:
    """Strip assessor notes, parenthetical examples and quotes from an indicator text."""
    if not text:
        return text
    cleaned = text

    for pattern in _PREFIX_PATTERNS:
        cleaned = pattern.sub(" ", cleaned)

    # Remove balanced parenthetical groups (examples / assessor hints), repeat for nesting.
    for _ in range(5):
        replaced = re.sub(r"\([^()]*\)", " ", cleaned)
        if replaced == cleaned:
            break
        cleaned = replaced
    # Drop a stray unbalanced "(" and everything after it.
    if "(" in cleaned:
        cleaned = cleaned[: cleaned.index("(")]

    # Remove balanced quoted fragments, then cut a stray unbalanced quote to the end.
    cleaned = re.sub(r"«[^»]*»", " ", cleaned)
    cleaned = re.sub(r'"[^"]*"', " ", cleaned)
    for quote in ('"', "«", "»"):
        if quote in cleaned:
            cleaned = cleaned[: cleaned.index(quote)]

    for pattern in _CUT_TO_END_PATTERNS:
        cleaned = pattern.sub("", cleaned)
    for pattern in _TOKEN_PATTERNS:
        cleaned = pattern.sub(" ", cleaned)

    # Drop leftover stray parentheses / plus signs.
    cleaned = re.sub(r"[()+]", " ", cleaned)
    cleaned = re.sub(r"\s*/\s*", " / ", cleaned)
    cleaned = re.sub(r"\.{2,}", ".", cleaned)
    cleaned = re.sub(r"\s+([,.;:])", r"\1", cleaned)
    cleaned = re.sub(r"[;,]\s*$", "", cleaned)
    cleaned = re.sub(r"\s{2,}", " ", cleaned)
    cleaned = cleaned.strip(" ;,-—/\t\n.")

    return cleaned if cleaned else text.strip()


def _format_course(course: Any) -> str:
    if isinstance(course, dict):
        title = str(course.get("title") or "").strip()
        url = str(course.get("url") or "").strip()
        if title and url:
            return f"{title} — {url}"
        return title or url
    return str(course)


def _recommendations(growth_zones: list[str]) -> list[str]:
    if not growth_zones:
        return ["Продолжать закреплять проявленные поведенческие индикаторы на рабочих задачах."]
    return [
        f"Отработать поведение: {item}. Подготовить 1-2 рабочих ситуации и заранее определить критерии успешного действия."
        for item in growth_zones[:6]
    ]


def _bullet_lines(items: list[str]) -> list[str]:
    if not items:
        return ["- Не выявлено по обработанным упражнениям."]
    return [f"- {item}" for item in items]


def _setup_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10)


def _add_heading(doc: Document, text: str, level: int = 1) -> None:
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(15 if level == 1 else 12)
    run.font.color.rgb = RGBColor(31, 78, 121 if level == 1 else 31)
    paragraph.paragraph_format.space_before = Pt(12)
    paragraph.paragraph_format.space_after = Pt(5)


def _add_bullet(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.add_run(str(text))


def _add_list_or_empty(doc: Document, items: list[str], empty_text: str) -> None:
    if not items:
        doc.add_paragraph(empty_text)
        return
    for item in items:
        _add_bullet(doc, item)


def _nearest_level_description(level: float) -> str:
    nearest = min(LEVEL_DESCRIPTIONS, key=lambda candidate: abs(candidate - level))
    return LEVEL_DESCRIPTIONS[nearest]
